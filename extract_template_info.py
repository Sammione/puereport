from docx import Document
import re
import json

def extract_info(path):
    doc = Document(path)
    info = {
        "placeholders": [],
        "tables": []
    }
    
    # Extract placeholders from paragraphs
    for p in doc.paragraphs:
        matches = re.findall(r'\[([^\]]+)\]', p.text)
        for m in matches:
            if m not in info["placeholders"]:
                info["placeholders"].append(m)
                
    # Extract table info
    for i, table in enumerate(doc.tables):
        headers = []
        if len(table.rows) > 0:
            headers = [cell.text.strip() for cell in table.rows[0].cells]
            
        # Check for placeholders in the table itself
        table_placeholders = []
        for row in table.rows:
            for cell in row.cells:
                matches = re.findall(r'\[([^\]]+)\]', cell.text)
                for m in matches:
                    if m not in table_placeholders:
                        table_placeholders.append(m)
        
        info["tables"].append({
            "index": i,
            "headers": headers,
            "placeholders": table_placeholders,
            "row_count": len(table.rows),
            "col_count": len(table.columns)
        })
        
    return info

if __name__ == "__main__":
    data = extract_info('PUE report templat- Copy.docx')
    with open('template_info.json', 'w') as f:
        json.dump(data, f, indent=2)
    print("Extracted info for", len(data["tables"]), "tables and", len(data["placeholders"]), "placeholders.")
