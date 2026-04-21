
import pandas as pd
from docx import Document
import re

def inspect_excel():
    print("--- Inspecting Excel: sheet 1.xlsx ---")
    try:
        df = pd.read_excel('sheet 1.xlsx')
        print("Columns:", list(df.columns))
        print("Shape:", df.shape)
        print("First row data:", df.iloc[0].to_dict())
    except Exception as e:
        print(f"Error reading Excel: {e}")

def inspect_docx():
    print("\n--- Inspecting Template: PUE report templat- Copy.docx ---")
    try:
        doc = Document('PUE report templat- Copy.docx')
        
        # Collect all text to search for potential placeholders
        full_text = []
        for p in doc.paragraphs:
            full_text.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        
        text_str = '\n'.join(full_text)
        
        # Look for patterns that look like placeholders e.g. {{...}}, [...], or just capitalized words known from excel
        print("First 10 paragraphs/items:")
        for t in full_text[:10]:
            if t.strip():
                print(f"- {t.strip()}")

        # Try to find specific usage of "Location" or common placeholders
        print("\nSearching for 'Location' or '[...]' patterns:")
        matches = re.findall(r'\[.*?\]|\{.*?\}|Location|Community', text_str, re.IGNORECASE)
        # Show unique matches (limited count)
        print(list(set(matches))[:20])

    except Exception as e:
        print(f"Error reading Docx: {e}")

if __name__ == "__main__":
    inspect_excel()
    inspect_docx()
