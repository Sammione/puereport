
from docx import Document
import sys

try:
    doc = Document('PUE report template Updated.docx')
    print("Paragraphs:")
    for p in doc.paragraphs[:20]: # Print first 20 paragraphs
        if p.text.strip():
            print(p.text)
    
    print("\nTables:")
    for len_table, table in enumerate(doc.tables):
        print(f"Table {len_table} has {len(table.rows)} rows and {len(table.columns)} columns")
        # Print first row text
        if table.rows:
            print([cell.text for cell in table.rows[0].cells])
except Exception as e:
    print(e)
