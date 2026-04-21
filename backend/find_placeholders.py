from docx import Document
import re

def inspect_doc(path):
    doc = Document(path)
    print(f"--- Document: {path} ---")
    print(f"Paragraphs count: {len(doc.paragraphs)}")
    print(f"Tables count: {len(doc.tables)}")
    
    print("\nFirst 10 paragraphs text:")
    for i, p in enumerate(doc.paragraphs[:10]):
        if p.text.strip():
            print(f"{i}: {p.text}")
            
    print("\nChecking for anything in brackets [...] or similar:")
    placeholders = set()
    for p in doc.paragraphs:
        matches = re.findall(r'\[([^\]]+)\]', p.text)
        for m in matches:
            placeholders.add(m)
            
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                matches = re.findall(r'\[([^\]]+)\]', cell.text)
                for m in matches:
                    placeholders.add(m)
                    
    print(f"Found {len(placeholders)} unique placeholders:")
    for p in sorted(list(placeholders)):
        print(f"  - {p}")

inspect_doc('PUE report template Updated.docx')
