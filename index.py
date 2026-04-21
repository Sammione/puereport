from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
import pandas as pd
import openai
from docx import Document
from docx.shared import Pt
import os
from dotenv import load_dotenv
import io
import json
import re
import uuid
from pydantic import BaseModel

load_dotenv()

session_mappings = {}

openai.api_key = os.getenv("OPENAI_API_KEY")

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Robust path finding
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_FILENAME = "PUE report templat- Copy.docx"

# Try looking in ../documents/ first (local dev structure), 
# then fallback to the current directory (repo root structure)
TEMPLATE_PATH = os.path.join(os.path.dirname(BASE_DIR), "documents", TEMPLATE_FILENAME)
if not os.path.exists(TEMPLATE_PATH):
    TEMPLATE_PATH = os.path.join(BASE_DIR, TEMPLATE_FILENAME)

def fill_table(table, data_rows):
    """
    Fills a table after clearing all rows except the header.
    """
    if not table.rows:
        return

    # Extract current header text to compare
    header_cells = [cell.text.strip().lower() for cell in table.rows[0].cells]
    
    # Delete all rows except the header
    while len(table.rows) > 1:
        tbl = table._tbl
        tr = table.rows[-1]._tr
        tbl.remove(tr)

    if not data_rows:
        return
    
    # Check if first data row is a header repeat
    start_idx = 0
    if len(data_rows) > 0:
        first_row_lower = [str(cell).strip().lower() for cell in data_rows[0]]
        # If substantially the same as header, skip it
        matches = sum(1 for i, val in enumerate(first_row_lower) if i < len(header_cells) and val == header_cells[i])
        if matches >= len(header_cells) * 0.8: # 80% match
            start_idx = 1

    for row_data in data_rows[start_idx:]:
        new_row = table.add_row()
        # Fill cells
        for j, val in enumerate(row_data):
            if j < len(new_row.cells):
                # Ensure it's a string and handle N/A
                text_val = str(val).strip() if (val is not None and str(val).strip().lower() != 'nan') else "N/A"
                new_row.cells[j].text = text_val

def docx_replace_placeholders(doc, replacements):
    """
    Replace text placeholders [KEY] in the doc.
    """
    # Create a normalized mapping
    norm_replacements = {str(k).upper().strip(): v for k, v in replacements.items()}
    
    def process_text(text):
        if not text: return text
        matches = re.findall(r'\[([^\]]+)\]', text)
        for match in matches:
            match_upper = match.upper().strip()
            if match_upper in norm_replacements:
                val = str(norm_replacements[match_upper])
                text = text.replace(f"[{match}]", val)
            else:
                text = text.replace(f"[{match}]", "N/A")
        return text

    for p in doc.paragraphs:
        p.text = process_text(p.text)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.text = process_text(p.text)

def apply_mapping_to_template(mapping):
    doc = Document(TEMPLATE_PATH)
    
    placeholders = mapping.get("placeholders", {})
    new_name = str(placeholders.get("COMMUNITY NAME") or placeholders.get("Community Name") or "Report")
    
    global_replacements = mapping.get("global_replacements", {})
    paragraph_updates = mapping.get("paragraph_updates", {}) 

    # 1. Apply paragraph-specific updates (indexed)
    for i, p in enumerate(doc.paragraphs):
        idx_str = str(i)
        if idx_str in paragraph_updates:
            p.text = paragraph_updates[idx_str]
    
    # 2. Apply global text replacements (thoroughly)
    def apply_global(text):
        if not text: return text
        for old_txt, new_txt in global_replacements.items():
            if old_txt and old_txt in text:
                text = text.replace(old_txt, str(new_txt))
        return text

    for p in doc.paragraphs:
        p.text = apply_global(p.text)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # Apply to cell text directly (handled via paragraphs)
                for p in cell.paragraphs:
                    p.text = apply_global(p.text)

    # 3. Standard placeholder replacement
    docx_replace_placeholders(doc, placeholders)

    # 4. Fill tables
    tables_data = mapping.get("tables", {})
    for i, table in enumerate(doc.tables):
        i_str = str(i)
        if i_str in tables_data:
            fill_table(table, tables_data[i_str])
        else:
            # If table is empty in mapping but has more than just header, clear it
            if len(table.rows) > 1:
                fill_table(table, [])

    clean_name = re.sub(r'[\\/*?:"<>|]', "", new_name).strip()
    output_name = f"Report_{clean_name}.docx"
    
    # Use a persistent temp directory in the root for local dev, 
    # but fall back to /tmp for Vercel/serverless environments.
    temp_dir = os.path.join(os.path.dirname(BASE_DIR), "temp")
    if not os.path.exists(temp_dir):
        try:
            os.makedirs(temp_dir)
        except:
            temp_dir = "/tmp"  # Fallback for serverless
            
    output_path = os.path.join(temp_dir, output_name)
    doc.save(output_path)
    return output_path, output_name

class RevisionRequest(BaseModel):
    session_id: str
    instruction: str

@app.post("/api/revise_report")
async def revise_report(req: RevisionRequest):
    try:
        if req.session_id not in session_mappings:
            raise HTTPException(status_code=404, detail="Session not found")
            
        current_mapping = session_mappings[req.session_id]
        
        prompt = f"""
You are a senior report engine. 
CURRENT JSON MAPPING OF THE REPORT:
{json.dumps(current_mapping)}

USER INSTRUCTION:
"{req.instruction}"

Return the FULL, UPDATED JSON mapping incorporating the user's modifications.
You must return the EXACT same structure (including 'placeholders', 'global_replacements', 'paragraph_updates', and 'tables'), just editing the requested values.
Keep the rest unchanged!
OUTPUT FORMAT MUST BE VALID JSON ONLY.
"""
        response = openai.chat.completions.create(
            model="gpt-4o-mini",
            response_format={"type": "json_object"},
            messages=[
                {"role": "system", "content": "You return JSON mapping matching the provided structure."},
                {"role": "user", "content": prompt}
            ],
            temperature=0,
            max_tokens=4096
        )
        
        new_mapping = json.loads(response.choices[0].message.content)
        session_mappings[req.session_id] = new_mapping
        
        output_path, output_name = apply_mapping_to_template(new_mapping)
        return FileResponse(output_path, media_type='application/vnd.wordprocessingml.document', filename=output_name)
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/generate_report")
async def generate_report(file: UploadFile = File(...)):
    try:
        print(f"Received file: {file.filename}")
        contents = await file.read()
        
        # Read Excel
        try:
            # Smart find header - check first 30 rows
            best_header_idx = 0
            max_keywords = -1
            
            temp_df_raw = pd.read_excel(io.BytesIO(contents), header=None, nrows=30)
            for idx, row in temp_df_raw.iterrows():
                row_str = " ".join([str(x).lower() for x in row if pd.notna(x)])
                # Expanded keyword list for better flexible matching
                keywords = [
                    'community', 'state', 'lga', 'name', 'coordinate', 'latitude', 'longitude',
                    'people interviewed', 'gender', 'crop', 'populati', 'household', 'location',
                    'site', 'village', 'town', 'area', 'region', 'gps', 'farmers', 'processor', 
                    'sme', 'mobility', 'energy', 'power', 'category', 'details'
                ]
                match_count = sum(1 for k in keywords if k in row_str)
                if match_count > max_keywords:
                    max_keywords = match_count
                    best_header_idx = idx
            
            print(f"Detected header at row {best_header_idx} with score {max_keywords}")
            df = pd.read_excel(io.BytesIO(contents), header=best_header_idx)
            df = df.dropna(how='all', axis=1).fillna("N/A")
            
            # GENERATE COMPREHENSIVE DATA SUMMARY FOR AI
            summary = []
            summary.append(f"### FULL EXCEL DATA SUMMARY (Total Rows: {len(df)})")
            summary.append(f"Columns Found: {', '.join(df.columns)}")
            
            # Add categorical summary (Counts)
            for col in df.columns:
                col_lower = str(col).lower()
                # Use a broader set of keywords to capture more relevant columns
                if any(x in col_lower for x in ['gender', 'ethnic', 'crop', 'activity', 'source', 'machine', 'income', 'people', 'business', 'type', 'category', 'status', 'use']):
                    counts = df[col].astype(str).value_counts().head(20).to_dict()
                    summary.append(f"Counts for '{col}': {json.dumps(counts)}")
                elif any(x in col_lower for x in ['population', 'household', 'total', 'count', 'number', 'qty', 'quantity', 'amount', 'sum', 'kilo', 'ton']):
                    # Summable numeric
                    try:
                        num_sum = pd.to_numeric(df[col], errors='coerce').sum()
                        summary.append(f"Total sum for '{col}': {num_sum}")
                    except:
                        pass
            
            # Send sample rows (representative)
            data_context = df.head(50).to_csv(index=False)
            if len(df) > 70:
                data_context += "\n\n... (Rows Skipped) ...\n\n" + df.tail(20).to_csv(index=False)
            
            # 1. Extract Ground Truth from first data row
            ground_truth = {}
            if len(df) > 0:
                first_row = df.iloc[0]
                # Look for name, lga, state in columns
                for col in df.columns:
                    c_low = str(col).lower()
                    if 'community' in c_low or 'name' in c_low: ground_truth['name'] = str(first_row.get(col, "N/A"))
                    if 'lga' in c_low: ground_truth['lga'] = str(first_row.get(col, "N/A"))
                    if 'state' in c_low: ground_truth['state'] = str(first_row.get(col, "N/A"))
            
            final_data_summary = "\n".join(summary) + "\n\n### REPRESENTATIVE DATA SAMPLE:\n" + data_context
            if ground_truth:
                final_data_summary = f"### DETECTED COMMUNITY GROUND TRUTH:\n{json.dumps(ground_truth, indent=2)}\n\n" + final_data_summary
            
        except Exception as e:
            import traceback
            traceback.print_exc()
            raise HTTPException(status_code=400, detail=f"Excel Parsing Error: {str(e)}")

        # Extract Template Info
        doc = Document(TEMPLATE_PATH)
        
        # 1. Extract Placeholders [...]
        detected_placeholders = set()
        for p in doc.paragraphs:
            matches = re.findall(r'\[([^\]]+)\]', p.text)
            for m in matches: detected_placeholders.add(m)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    matches = re.findall(r'\[([^\]]+)\]', cell.text)
                    for m in matches: detected_placeholders.add(m)

        # 2. Extract Table Context
        table_context = []
        for i, table in enumerate(doc.tables):
            headers = [cell.text.strip() for cell in table.rows[0].cells] if table.rows else []
            table_context.append({"index": i, "headers": headers})

        # 3. Extract Paragraph Context (Search for community names or specific data to replace)
        paragraph_context = []
        potential_stale_names = ["Agadagba", "Lagos", "Community Name"] # Common ones to watch
        
        for i, p in enumerate(doc.paragraphs):
            text = p.text.strip()
            # If paragraph contains likely stale data or digits
            if len(text) > 20 and (any(char.isdigit() for char in text) or any(name.lower() in text.lower() for name in potential_stale_names) or "Overview" in text or "Background" in text):
                paragraph_context.append({"index": i, "text": text})

        # Prompt AI for JSON mapping
        prompt = """
You are a senior report engine. You must convert field survey data into a professional PUE report mapping.

### THE DATA FROM EXCEL:
\"\"\"
{DATA_SUMMARY}
\"\"\"

### DOCUMENT PLACEHOLDERS FOUND:
{PLACEHOLDERS}

### REPORT TABLES STRUCTURE:
{TABLES}

### POTENTIAL STALE PARAGRAPHS IN TEMPLATE:
{PARAGRAPHS}

### RIGID INSTRUCTIONS:
1. **Dynamic Replacement**: The template contains stale data from a previous community (e.g., "Agadagba", specific kWp values, etc.). You must identify these and provide replacements in 'global_replacements'.
2. **Factual Community Overview**:
   - The 'Community Overview' section in the template is about a different place (Agadagba).
   - **DO NOT** just swap "Agadagba" with the new name.
   - **IF** you recognize the new community as a real geographical location (e.g., Lagos, Nigeria), use your internal knowledge to provide a concise, ACCURATE description of its geography, people, and economy.
   - **IF** it is a small rural village, use facts from the Excel (population, economic activities) to synthesize a fresh description.
   - **MANDATORY**: Remove references to specific rulers (Obas), ethnic groups (e.g. Ikale), or history that belongs to the old community.
3. **Placeholders**: Fill all placeholders found in the document. 
4. **Table Data**: 
   - Fill all professional tables with data from the Excel summary.
   - **IMPORTANT**: DO NOT include the header row in the table data JSON. The engine already has the headers. Start directly with the data rows.
5. **Community Names**: Ensure EVERY mention of "Agadagba" is replaced with the new community name in 'global_replacements'.
6. **Consistency**: Use consistent naming. Use the detected Ground Truth if provided.

### OUTPUT FORMAT:
{
  "placeholders": {
    "COMMUNITY NAME": "...",
    "Date": "...",
    ...
  },
  "global_replacements": {
     "Agadagba": "New Community Name",
     "Old Value": "New Value"
  },
  "paragraph_updates": {
     "idx": "New full text for this paragraph..."
  },
  "tables": {
     "0": [["Row 1 Col 1", "Col 2", ...]],
     ...
  }
}
""".replace("{DATA_SUMMARY}", final_data_summary)\
   .replace("{PLACEHOLDERS}", json.dumps(list(detected_placeholders)))\
   .replace("{TABLES}", json.dumps(table_context[:65], indent=2))\
   .replace("{PARAGRAPHS}", json.dumps(paragraph_context[:60], indent=2))

        print("Requesting AI for mapping...")
        response = openai.chat.completions.create(
            model="gpt-4o-mini",
            response_format={"type": "json_object"},
            messages=[
                {"role": "system", "content": "You are a professional report analyst. You extract and return survey data in JSON format only. You provide factual, accurate geographic descriptions for community overviews, avoiding blind template substitution."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.2,
            max_tokens=4096
        )
        
        mapping = json.loads(response.choices[0].message.content)
        print("Mapping received.")

        session_id = str(uuid.uuid4())
        session_mappings[session_id] = mapping
        
        output_path, output_name = apply_mapping_to_template(mapping)
        
        return FileResponse(output_path, media_type='application/vnd.wordprocessingml.document', filename=output_name, headers={"X-Session-ID": session_id, "Access-Control-Expose-Headers": "X-Session-ID"})

    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8700)
